"""
DOCX parser for GDPdU/GoBD dossier documents.

Extracts paragraphs and tables from Word documents, preserving heading
structure for context. Each paragraph or table row becomes a NormalizedRecord
of type document_text.
"""

from __future__ import annotations

import logging
import re
import uuid
from pathlib import Path

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from app.normalization.models import (
    EntityRef,
    NormalizedRecord,
    RecordType,
    SourceProvenance,
)

logger = logging.getLogger(__name__)

NAMESPACE_URL = uuid.NAMESPACE_URL

# Patterns for entity extraction from text
_EUR_AMOUNT_RE = re.compile(
    r"(?:EUR|€)\s*([+-]?\d{1,3}(?:[.,]\d{3})*(?:[.,]\d{1,2})?)"
    r"|([+-]?\d{1,3}(?:[.,]\d{3})*(?:[.,]\d{1,2})?)\s*(?:EUR|€)",
    re.IGNORECASE,
)
_USER_ID_RE = re.compile(r"\b(MV-U\d{2,3})\b")
_ACCOUNT_NUMBER_RE = re.compile(r"\b(\d{4,6})\b")


def _make_record_id(dossier_id: str, file_id: str, row: int, index: int) -> str:
    """Generate stable UUID5 for a record."""
    dossier_ns = uuid.uuid5(NAMESPACE_URL, f"dossier:{dossier_id}")
    return str(uuid.uuid5(dossier_ns, f"{file_id}:{row}:{index}"))


def _extract_entities_from_text(text: str) -> list[EntityRef]:
    """Extract entity references from document text."""
    entities: list[EntityRef] = []
    seen: set[tuple[str, str]] = set()

    # Extract user IDs (e.g., MV-U05)
    for match in _USER_ID_RE.finditer(text):
        user_id = match.group(1)
        key = ("user", user_id)
        if key not in seen:
            seen.add(key)
            entities.append(
                EntityRef(entity_type="user", entity_id=user_id, label=user_id)
            )

    return entities


def _extract_amount_from_text(text: str) -> float | None:
    """Extract the first EUR amount from text."""
    match = _EUR_AMOUNT_RE.search(text)
    if match:
        raw = match.group(1) or match.group(2)
        if raw:
            # Handle German number format
            cleaned = raw.strip()
            if "," in cleaned and "." in cleaned:
                # 1.234,56 format
                cleaned = cleaned.replace(".", "").replace(",", ".")
            elif "," in cleaned:
                cleaned = cleaned.replace(",", ".")
            try:
                return float(cleaned)
            except ValueError:
                pass
    return None


def parse_docx_file(
    file_path: Path,
    relative_path: str,
    dossier_id: str,
    file_id: str,
) -> list[NormalizedRecord]:
    """
    Parse a DOCX file into NormalizedRecords.

    Extracts paragraphs and tables. Each paragraph becomes a document_text
    record. Table rows are extracted as separate records. Heading structure
    is preserved as context in the text_content.

    Args:
        file_path: Absolute path to the DOCX file.
        relative_path: Path relative to dossier root.
        dossier_id: ID of the parent dossier.
        file_id: Stable ID of this file in the manifest.

    Returns:
        List of NormalizedRecords. Empty list on failure.
    """
    records: list[NormalizedRecord] = []

    try:
        doc = Document(str(file_path))
    except (PackageNotFoundError, Exception) as exc:
        logger.warning(
            "Failed to open DOCX file %s: %s", relative_path, exc
        )
        return []

    # Track heading context
    current_heading: str = ""
    paragraph_index: int = 0

    # Process paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            paragraph_index += 1
            continue

        # Track headings for context
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            current_heading = text

        # Build text with heading context
        if current_heading and text != current_heading:
            full_text = f"[{current_heading}] {text}"
        else:
            full_text = text

        try:
            entities = _extract_entities_from_text(text)
            amount = _extract_amount_from_text(text)

            relationships: dict[str, str] = {}
            for entity in entities:
                relationships[entity.entity_type] = entity.entity_id

            record_id = _make_record_id(
                dossier_id, file_id, paragraph_index, 0
            )

            record = NormalizedRecord(
                record_id=record_id,
                dossier_id=dossier_id,
                record_type=RecordType.document_text,
                source=SourceProvenance(
                    file_id=file_id,
                    relative_path=relative_path,
                    paragraph=paragraph_index,
                ),
                entities=entities,
                relationships=relationships,
                data={
                    "heading": current_heading if current_heading else None,
                    "style": para.style.name if para.style else None,
                },
                text_content=full_text,
                amount=amount,
                currency="EUR" if amount is not None else None,
            )
            records.append(record)

        except Exception as exc:
            logger.warning(
                "Skipping paragraph %d in %s: %s",
                paragraph_index,
                relative_path,
                exc,
            )

        paragraph_index += 1

    # Process tables
    for table_idx, table in enumerate(doc.tables):
        # Extract header row
        if not table.rows:
            continue

        header_cells = [
            cell.text.strip() for cell in table.rows[0].cells
        ]

        for row_idx, row in enumerate(table.rows[1:], start=1):
            try:
                cells = [cell.text.strip() for cell in row.cells]

                # Skip empty rows
                if all(not c for c in cells):
                    continue

                # Build row text
                row_parts: list[str] = []
                row_data: dict[str, str | int | float | None] = {}
                for col_idx, cell_text in enumerate(cells):
                    if col_idx < len(header_cells) and header_cells[col_idx]:
                        col_name = header_cells[col_idx]
                    else:
                        col_name = f"column_{col_idx}"
                    row_data[col_name] = cell_text if cell_text else None
                    if cell_text:
                        row_parts.append(f"{col_name}: {cell_text}")

                row_text = "; ".join(row_parts)
                full_row_text = (
                    f"[Table {table_idx + 1}, Row {row_idx}] {row_text}"
                )

                entities = _extract_entities_from_text(row_text)
                amount = _extract_amount_from_text(row_text)

                relationships: dict[str, str] = {}
                for entity in entities:
                    relationships[entity.entity_type] = entity.entity_id

                # Use paragraph_index + table offset for unique row identifiers
                provenance_row = paragraph_index + table_idx * 1000 + row_idx
                record_id = _make_record_id(
                    dossier_id, file_id, provenance_row, 0
                )

                record = NormalizedRecord(
                    record_id=record_id,
                    dossier_id=dossier_id,
                    record_type=RecordType.document_text,
                    source=SourceProvenance(
                        file_id=file_id,
                        relative_path=relative_path,
                        row_number=row_idx,
                        columns=header_cells if header_cells else None,
                    ),
                    entities=entities,
                    relationships=relationships,
                    data=row_data,
                    text_content=full_row_text,
                    amount=amount,
                    currency="EUR" if amount is not None else None,
                )
                records.append(record)

            except Exception as exc:
                logger.warning(
                    "Skipping table %d row %d in %s: %s",
                    table_idx,
                    row_idx,
                    relative_path,
                    exc,
                )
                continue

    logger.info(
        "Parsed %d records from DOCX %s", len(records), relative_path
    )
    return records
